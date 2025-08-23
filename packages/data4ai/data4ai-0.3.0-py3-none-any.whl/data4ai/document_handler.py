"""Document file operations for Data4AI - Extract and process PDF, DOCX, MD, and TXT files."""

import logging
import re
from pathlib import Path
from typing import Any, Optional, Union

from data4ai.exceptions import ValidationError

logger = logging.getLogger("data4ai")

# Check for optional dependencies
try:
    import pypdf  # noqa: F401

    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import pdfplumber  # noqa: F401

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown  # noqa: F401

    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


class DocumentHandler:
    """Handle document file operations for dataset generation."""

    @staticmethod
    def detect_document_type(file_path: Path) -> str:
        """Detect document type from file extension.

        Args:
            file_path: Path to document file

        Returns:
            Document type (pdf, docx, md, txt)

        Raises:
            ValidationError: If file type is not supported
        """
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return "pdf"
        elif suffix in [".docx", ".doc"]:
            return "docx"
        elif suffix in [".md", ".markdown"]:
            return "md"
        elif suffix in [".txt", ".text"]:
            return "txt"
        else:
            raise ValidationError(
                f"Unsupported document type: {suffix}. "
                "Supported types: .pdf, .docx, .md, .txt"
            )

    @staticmethod
    def extract_text(file_path: Path, use_advanced: bool = False) -> str:
        """Extract text from document file.

        Args:
            file_path: Path to document file
            use_advanced: Use advanced extraction (pdfplumber for PDFs)

        Returns:
            Extracted text content

        Raises:
            ValidationError: If extraction fails or dependencies missing
        """
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        doc_type = DocumentHandler.detect_document_type(file_path)

        if doc_type == "pdf":
            return DocumentHandler._extract_pdf_text(file_path, use_advanced)
        elif doc_type == "docx":
            return DocumentHandler._extract_docx_text(file_path)
        elif doc_type == "md":
            return DocumentHandler._extract_markdown_text(file_path)
        elif doc_type == "txt":
            return DocumentHandler._extract_txt_text(file_path)
        else:
            raise ValidationError(f"Unsupported document type: {doc_type}")

    @staticmethod
    def _extract_pdf_text(file_path: Path, use_advanced: bool = False) -> str:
        """Extract text from PDF file.

        Args:
            file_path: Path to PDF file
            use_advanced: Use pdfplumber for better extraction

        Returns:
            Extracted text
        """
        if use_advanced and PDFPLUMBER_AVAILABLE:
            logger.info("Using pdfplumber for advanced PDF extraction")
            try:
                import pdfplumber

                text_parts = []

                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"--- Page {page_num} ---\n{page_text}")

                return "\n\n".join(text_parts)
            except Exception as e:
                logger.warning(
                    f"pdfplumber extraction failed: {e}, falling back to pypdf"
                )
                # Fall back to pypdf

        if not PYPDF_AVAILABLE:
            raise ValidationError(
                "PDF support not available. Please install with: "
                "pip install data4ai[docs] or pip install pypdf"
            )

        try:
            import pypdf

            reader = pypdf.PdfReader(file_path)
            text_parts = []

            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")

            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValidationError(f"Failed to extract PDF text: {str(e)}") from e

    @staticmethod
    def _extract_docx_text(file_path: Path) -> str:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            raise ValidationError(
                "DOCX support not available. Please install with: "
                "pip install data4ai[docs] or pip install python-docx"
            )

        try:
            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            return "\n\n".join(paragraphs)
        except Exception as e:
            raise ValidationError(f"Failed to extract DOCX text: {str(e)}") from e

    @staticmethod
    def _extract_markdown_text(file_path: Path) -> str:
        """Extract text from Markdown file.

        Args:
            file_path: Path to Markdown file

        Returns:
            Plain text (with markdown stripped if library available)
        """
        try:
            with open(file_path, encoding="utf-8") as f:
                content = f.read()

            if MARKDOWN_AVAILABLE:
                # Convert markdown to HTML then strip tags
                from html.parser import HTMLParser

                import markdown

                class HTMLStripper(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.reset()
                        self.strict = False
                        self.convert_charrefs = True
                        self.text = []

                    def handle_data(self, data):
                        self.text.append(data)

                    def get_text(self):
                        return "".join(self.text)

                html = markdown.markdown(content)
                stripper = HTMLStripper()
                stripper.feed(html)
                return stripper.get_text()
            else:
                # Basic markdown stripping
                # Remove code blocks
                content = re.sub(r"```.*?```", "", content, flags=re.DOTALL)
                content = re.sub(r"`[^`]+`", "", content)
                # Remove headers
                content = re.sub(r"^#+\s+", "", content, flags=re.MULTILINE)
                # Remove emphasis
                content = re.sub(r"[*_]{1,2}([^*_]+)[*_]{1,2}", r"\1", content)
                # Remove links
                content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)
                # Remove images
                content = re.sub(r"!\[([^\]]*)\]\([^)]+\)", "", content)

                return content
        except Exception as e:
            raise ValidationError(f"Failed to extract Markdown text: {str(e)}") from e

    @staticmethod
    def _extract_txt_text(file_path: Path) -> str:
        """Extract text from TXT file.

        Args:
            file_path: Path to TXT file

        Returns:
            File content
        """
        try:
            with open(file_path, encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, encoding="latin-1") as f:
                    return f.read()
            except Exception as e:
                raise ValidationError(f"Failed to extract text: {str(e)}") from e
        except Exception as e:
            raise ValidationError(f"Failed to extract text: {str(e)}") from e

    @staticmethod
    def extract_chunks(
        file_path: Path,
        chunk_size: int = 1000,
        overlap: int = 200,
        use_advanced: bool = False,
    ) -> list[dict[str, Any]]:
        """Extract text in chunks for processing.

        Args:
            file_path: Path to document file
            chunk_size: Size of each chunk in characters
            overlap: Overlap between chunks for context
            use_advanced: Use advanced extraction

        Returns:
            List of chunks with metadata
        """
        text = DocumentHandler.extract_text(file_path, use_advanced)

        if not text:
            return []

        # Clamp overlap to avoid infinite/very slow progression
        if chunk_size <= 0:
            chunk_size = 1000
        if overlap < 0:
            overlap = 0
        if overlap >= chunk_size:
            overlap = max(0, chunk_size - 1)

        chunks = []
        start = 0
        chunk_id = 0

        while start < len(text):
            end = min(start + chunk_size, len(text))

            # Try to find a good break point (sentence end)
            if end < len(text):
                # Look for sentence endings
                for sep in [".\n", ". ", "!\n", "! ", "?\n", "? "]:
                    last_sep = text.rfind(sep, start, end)
                    if last_sep != -1:
                        end = last_sep + len(sep)
                        break

            chunk_text = text[start:end].strip()

            if chunk_text:
                chunks.append(
                    {
                        "id": chunk_id,
                        "text": chunk_text,
                        "start": start,
                        "end": end,
                        "source": file_path.name,
                    }
                )
                chunk_id += 1

            # Move start position with overlap and ensure forward progress
            start = max(start + 1, end - overlap) if end < len(text) else end

        logger.info(f"Extracted {len(chunks)} chunks from {file_path.name}")
        return chunks

    @staticmethod
    def is_supported_file(file_path: Path) -> bool:
        """Check if file is a supported document type.

        Args:
            file_path: Path to check

        Returns:
            True if file is supported
        """
        try:
            DocumentHandler.detect_document_type(file_path)
            return True
        except ValidationError:
            return False

    @staticmethod
    def scan_folder(
        folder_path: Path,
        recursive: bool = True,
        file_types: Optional[list[str]] = None,
    ) -> list[Path]:
        """Scan folder for supported documents.

        Args:
            folder_path: Path to folder
            recursive: Whether to scan subfolders
            file_types: Specific file types to include (pdf, docx, md, txt)

        Returns:
            List of document paths
        """
        if not folder_path.exists():
            raise FileNotFoundError(f"Folder not found: {folder_path}")

        if not folder_path.is_dir():
            raise ValidationError(f"Path is not a directory: {folder_path}")

        # Default to all supported types
        if file_types is None:
            file_types = ["pdf", "docx", "doc", "md", "markdown", "txt", "text"]
        else:
            # Normalize file types
            file_types = [ft.lower() for ft in file_types]

        documents = []

        # Define supported extensions
        extensions = []
        if "pdf" in file_types:
            extensions.append("*.pdf")
        if "docx" in file_types or "doc" in file_types:
            extensions.extend(["*.docx", "*.doc"])
        if "md" in file_types or "markdown" in file_types:
            extensions.extend(["*.md", "*.markdown"])
        if "txt" in file_types or "text" in file_types:
            extensions.extend(["*.txt", "*.text"])

        # Scan for files
        for ext in extensions:
            if recursive:
                documents.extend(folder_path.rglob(ext))
            else:
                documents.extend(folder_path.glob(ext))

        # Filter out hidden files and ensure unique paths
        documents = list(
            {doc for doc in documents if not doc.name.startswith(".") and doc.is_file()}
        )

        # Sort for consistent ordering
        documents.sort()

        logger.info(f"Found {len(documents)} documents in {folder_path}")
        return documents

    @staticmethod
    def extract_from_multiple(
        file_paths: list[Path], use_advanced: bool = False, combine: bool = True
    ) -> Union[str, dict[str, str]]:
        """Extract text from multiple documents.

        Args:
            file_paths: List of document paths
            use_advanced: Use advanced extraction
            combine: Whether to combine into single text or return dict

        Returns:
            Combined text string or dict mapping paths to text
        """
        texts = {}

        for file_path in file_paths:
            try:
                text = DocumentHandler.extract_text(file_path, use_advanced)
                if text:
                    texts[str(file_path)] = text
            except Exception as e:
                logger.warning(f"Failed to extract from {file_path}: {e}")
                continue

        if combine:
            # Combine all texts with document separators
            combined_parts = []
            for path, text in texts.items():
                doc_name = Path(path).name
                combined_parts.append(f"=== Document: {doc_name} ===\n\n{text}")
            return "\n\n".join(combined_parts)
        else:
            return texts

    @staticmethod
    def _analyze_folder_structure(
        root_folder: Path, file_paths: list[Path]
    ) -> dict[str, Any]:
        """Analyze folder structure to organize documents by subfolder.

        Args:
            root_folder: The root folder being scanned
            file_paths: List of document file paths found

        Returns:
            Dictionary with folder structure information
        """
        structure = {}

        # Group files by their relative subfolder
        for file_path in file_paths:
            try:
                # Get relative path from root
                relative_path = file_path.relative_to(root_folder)

                # Get the parent folder (subfolder)
                if len(relative_path.parts) > 1:
                    # File is in a subfolder
                    subfolder = relative_path.parts[0]
                    if subfolder not in structure:
                        structure[subfolder] = []
                    structure[subfolder].append(file_path)
                else:
                    # File is in root folder
                    if "root" not in structure:
                        structure["root"] = []
                    structure["root"].append(file_path)

            except ValueError:
                # File is not under root_folder, add to root
                if "root" not in structure:
                    structure["root"] = []
                structure["root"].append(file_path)

        return {
            "root_folder": str(root_folder),
            "subfolders": structure,
            "has_subfolders": len([k for k in structure if k != "root"]) > 0,
        }

    @staticmethod
    def extract_chunks_from_multiple(
        file_paths: list[Path],
        chunk_size: int = 1000,
        overlap: int = 200,
        use_advanced: bool = False,
    ) -> list[dict[str, Any]]:
        """Extract chunks from multiple documents.

        Args:
            file_paths: List of document paths
            chunk_size: Size of each chunk
            overlap: Overlap between chunks
            use_advanced: Use advanced extraction

        Returns:
            List of all chunks with metadata
        """
        all_chunks = []

        for file_path in file_paths:
            try:
                chunks = DocumentHandler.extract_chunks(
                    file_path,
                    chunk_size=chunk_size,
                    overlap=overlap,
                    use_advanced=use_advanced,
                )
                # Add file path to chunk metadata
                for chunk in chunks:
                    chunk["file_path"] = str(file_path)
                all_chunks.extend(chunks)
            except Exception as e:
                logger.warning(f"Failed to extract chunks from {file_path}: {e}")
                continue

        logger.info(
            f"Extracted {len(all_chunks)} total chunks from {len(file_paths)} documents"
        )
        return all_chunks

    @staticmethod
    def convert_pdf_to_markdown(
        pdf_path: Path,
        output_path: Optional[Path] = None,
        use_advanced: bool = True,
        clean_output: bool = True,
    ) -> Path:
        """Convert PDF to Markdown format.

        Args:
            pdf_path: Path to PDF file
            output_path: Output path for markdown file (defaults to same name with .md)
            use_advanced: Use pdfplumber for better extraction
            clean_output: Clean and format the output text

        Returns:
            Path to the created markdown file
        """
        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")

        if not pdf_path.suffix.lower() == ".pdf":
            raise ValidationError(f"Not a PDF file: {pdf_path}")

        # Extract text from PDF
        logger.info(f"Converting PDF to markdown: {pdf_path.name}")
        text = DocumentHandler._extract_pdf_text(pdf_path, use_advanced)

        if clean_output:
            # Clean up common PDF extraction issues
            import re

            # Remove excessive whitespace
            text = re.sub(r"\n{3,}", "\n\n", text)
            text = re.sub(r" {2,}", " ", text)

            # Fix common OCR issues
            text = re.sub(
                r"(?<=[a-z])(?=[A-Z])", " ", text
            )  # Add space between camelCase

            # Format page breaks
            text = re.sub(r"--- Page (\d+) ---", r"\n\n---\n\n## Page \1\n\n", text)

            # Clean up hyphenated words at line breaks
            text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)

        # Determine output path
        if output_path is None:
            output_path = pdf_path.with_suffix(".md")

        # Write markdown file
        with open(output_path, "w", encoding="utf-8") as f:
            # Add header
            f.write(f"# {pdf_path.stem}\n\n")
            f.write(f"*Converted from PDF: {pdf_path.name}*\n\n")
            f.write("---\n\n")
            f.write(text)

        logger.info(f"PDF converted to markdown: {output_path}")
        return output_path

    @staticmethod
    def convert_pdfs_in_folder(
        folder_path: Path,
        output_folder: Optional[Path] = None,
        recursive: bool = True,
        use_advanced: bool = True,
        delete_pdfs: bool = False,
    ) -> list[Path]:
        """Convert all PDFs in a folder to Markdown.

        Args:
            folder_path: Path to folder containing PDFs
            output_folder: Folder for markdown files (defaults to same folder)
            recursive: Process subfolders
            use_advanced: Use pdfplumber for better extraction
            delete_pdfs: Delete PDF files after conversion

        Returns:
            List of created markdown file paths
        """
        if not folder_path.exists():
            raise FileNotFoundError(f"Folder not found: {folder_path}")

        if not folder_path.is_dir():
            raise ValidationError(f"Not a directory: {folder_path}")

        # Find all PDF files
        pdf_files = []
        if recursive:
            pdf_files = list(folder_path.rglob("*.pdf"))
        else:
            pdf_files = list(folder_path.glob("*.pdf"))

        if not pdf_files:
            logger.info(f"No PDF files found in {folder_path}")
            return []

        logger.info(f"Found {len(pdf_files)} PDF files to convert")

        # Determine output folder
        if output_folder is None:
            output_folder = folder_path
        else:
            output_folder.mkdir(parents=True, exist_ok=True)

        # Convert each PDF
        converted_files = []
        for pdf_file in pdf_files:
            try:
                # Maintain folder structure if recursive
                if recursive and output_folder != folder_path:
                    relative_path = pdf_file.relative_to(folder_path)
                    output_path = output_folder / relative_path.with_suffix(".md")
                    output_path.parent.mkdir(parents=True, exist_ok=True)
                else:
                    output_path = output_folder / pdf_file.with_suffix(".md").name

                # Convert PDF to Markdown
                md_path = DocumentHandler.convert_pdf_to_markdown(
                    pdf_file, output_path, use_advanced=use_advanced
                )
                converted_files.append(md_path)

                # Delete PDF if requested
                if delete_pdfs:
                    pdf_file.unlink()
                    logger.info(f"Deleted PDF: {pdf_file}")

            except Exception as e:
                logger.error(f"Failed to convert {pdf_file}: {e}")
                continue

        logger.info(f"Successfully converted {len(converted_files)} PDF files")
        return converted_files

    @staticmethod
    def prepare_for_generation(
        input_path: Union[Path, list[Path]],
        extraction_type: str = "qa",
        chunk_size: int = 1000,
        use_advanced: bool = False,
        recursive: bool = True,
    ) -> dict[str, Any]:
        """Prepare document(s) for dataset generation.

        Args:
            input_path: Path to document file, folder, or list of paths
            extraction_type: Type of extraction (qa, summary, instruction)
            chunk_size: Size of chunks for processing
            use_advanced: Use advanced extraction
            recursive: For folders, whether to scan recursively

        Returns:
            Prepared data for generation
        """
        # Handle different input types
        if isinstance(input_path, list):
            # List of paths provided
            file_paths = input_path
            input_type = "multiple"
        elif isinstance(input_path, Path):
            if input_path.is_dir():
                # Folder provided - scan for documents
                file_paths = DocumentHandler.scan_folder(
                    input_path, recursive=recursive
                )
                input_type = "folder"
                if not file_paths:
                    raise ValidationError(
                        f"No supported documents found in {input_path}"
                    )
            else:
                # Single file provided
                file_paths = [input_path]
                input_type = "file"
        else:
            # Convert string to Path
            input_path = Path(input_path)
            return DocumentHandler.prepare_for_generation(
                input_path, extraction_type, chunk_size, use_advanced, recursive
            )

        # Extract chunks from all documents
        if len(file_paths) == 1:
            # Single document
            doc_type = DocumentHandler.detect_document_type(file_paths[0])
            chunks = DocumentHandler.extract_chunks(
                file_paths[0], chunk_size=chunk_size, use_advanced=use_advanced
            )

            return {
                "document_type": doc_type,
                "document_name": file_paths[0].name,
                "extraction_type": extraction_type,
                "chunks": chunks,
                "total_chunks": len(chunks),
                "total_documents": 1,
                "input_type": input_type,
                "metadata": {
                    "chunk_size": chunk_size,
                    "file_size": file_paths[0].stat().st_size,
                    "advanced_extraction": use_advanced,
                },
            }
        else:
            # Multiple documents
            chunks = DocumentHandler.extract_chunks_from_multiple(
                file_paths, chunk_size=chunk_size, use_advanced=use_advanced
            )

            # Get document types and sizes
            doc_types = set()
            total_size = 0
            for fp in file_paths:
                try:
                    doc_types.add(DocumentHandler.detect_document_type(fp))
                    total_size += fp.stat().st_size
                except Exception:
                    pass

            # Analyze folder structure if input was a folder
            folder_structure = None
            if input_type == "folder" and isinstance(input_path, Path):
                folder_structure = DocumentHandler._analyze_folder_structure(
                    input_path, file_paths
                )

            return {
                "document_type": "mixed" if len(doc_types) > 1 else list(doc_types)[0],
                "document_names": [fp.name for fp in file_paths],
                "extraction_type": extraction_type,
                "chunks": chunks,
                "total_chunks": len(chunks),
                "total_documents": len(file_paths),
                "input_type": input_type,
                "folder_structure": folder_structure,
                "metadata": {
                    "chunk_size": chunk_size,
                    "total_file_size": total_size,
                    "advanced_extraction": use_advanced,
                    "document_types": list(doc_types),
                },
            }
