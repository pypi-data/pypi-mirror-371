"""
File processing utilities for LLM requests.
"""

import json
import csv
import xml.etree.ElementTree as ET
from typing import Dict, List, Optional, Union, Any, Tuple
from pathlib import Path
import logging
import mimetypes
from dataclasses import dataclass

from ..models.request import ProcessedFile
from .token_estimator import TokenEstimator

logger = logging.getLogger(__name__)


@dataclass
class ProcessingResult:
    """Result of file processing operation."""
    success: bool
    content: str = ""
    content_type: str = ""
    token_count: int = 0
    processing_method: str = ""
    warnings: List[str] = None
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []
        if self.metadata is None:
            self.metadata = {}


class FileProcessor:
    """Processes various file types for LLM consumption."""
    
    # Supported file extensions and their MIME types
    SUPPORTED_EXTENSIONS = {
        # Text files
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.rst': 'text/x-rst',
        '.log': 'text/plain',
        
        # Code files
        '.py': 'text/x-python',
        '.js': 'text/javascript',
        '.ts': 'text/typescript',
        '.html': 'text/html',
        '.css': 'text/css',
        '.sql': 'text/x-sql',
        '.sh': 'text/x-shellscript',
        '.yaml': 'text/yaml',
        '.yml': 'text/yaml',
        
        # Data files
        '.json': 'application/json',
        '.csv': 'text/csv',
        '.xml': 'application/xml',
        '.tsv': 'text/tab-separated-values',
        
        # Document files
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.rtf': 'application/rtf',
        
        # Image files (for metadata extraction)
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.webp': 'image/webp',
        '.tiff': 'image/tiff',
        '.svg': 'image/svg+xml',
    }
    
    def __init__(self, 
                 token_estimator: Optional[TokenEstimator] = None,
                 max_file_size: int = 10 * 1024 * 1024,  # 10MB
                 encoding: str = 'utf-8'):
        """
        Initialize the file processor.
        
        Args:
            token_estimator: Token estimator for content analysis
            max_file_size: Maximum file size to process (bytes)
            encoding: Default text encoding
        """
        self.token_estimator = token_estimator or TokenEstimator()
        self.max_file_size = max_file_size
        self.encoding = encoding
    
    def process_file(self, 
                    file_path: Union[str, Path], 
                    model: str,
                    content_type: Optional[str] = None) -> ProcessingResult:
        """
        Process a file and extract its content.
        
        Args:
            file_path: Path to the file
            model: Model name for token estimation
            content_type: Optional content type override
            
        Returns:
            ProcessingResult with extracted content and metadata
        """
        file_path = Path(file_path)
        
        # Check if file exists
        if not file_path.exists():
            return ProcessingResult(
                success=False,
                processing_method="file_not_found",
                warnings=[f"File not found: {file_path}"]
            )
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            return ProcessingResult(
                success=False,
                processing_method="file_too_large",
                warnings=[f"File too large: {file_size} bytes (max: {self.max_file_size})"]
            )
        
        # Determine content type
        if not content_type:
            content_type = self._detect_content_type(file_path)
        
        # Process based on content type
        try:
            if content_type.startswith('text/') or content_type in ['application/json', 'application/xml']:
                result = self._process_text_file(file_path, content_type, model)
            elif content_type.startswith('image/'):
                result = self._process_image_file(file_path, content_type, model)
            elif content_type == 'application/pdf':
                result = self._process_pdf_file(file_path, content_type, model)
            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                result = self._process_word_file(file_path, content_type, model)
            elif content_type == 'application/rtf':
                result = self._process_rtf_file(file_path, content_type, model)
            else:
                result = ProcessingResult(
                    success=False,
                    content_type=content_type,
                    processing_method="unsupported_type",
                    warnings=[f"Unsupported content type: {content_type}"]
                )
            
            # Add file metadata
            result.metadata.update({
                "file_path": str(file_path),
                "file_size": file_size,
                "file_extension": file_path.suffix.lower()
            })
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return ProcessingResult(
                success=False,
                content_type=content_type,
                processing_method="processing_error",
                warnings=[f"Processing error: {str(e)}"]
            )
    
    def _detect_content_type(self, file_path: Path) -> str:
        """Detect content type from file extension and content."""
        # Try extension-based detection first
        extension = file_path.suffix.lower()
        if extension in self.SUPPORTED_EXTENSIONS:
            return self.SUPPORTED_EXTENSIONS[extension]
        
        # Try mimetypes library
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type:
            return mime_type
        
        # Default to plain text
        return 'text/plain'
    
    def _process_text_file(self, 
                          file_path: Path, 
                          content_type: str, 
                          model: str) -> ProcessingResult:
        """Process text-based files."""
        try:
            # Read file content
            with open(file_path, 'r', encoding=self.encoding, errors='replace') as f:
                raw_content = f.read()
            
            # Process based on specific content type
            if content_type == 'application/json':
                processed_content, method, warnings = self._process_json_content(raw_content)
            elif content_type == 'text/csv':
                processed_content, method, warnings = self._process_csv_content(raw_content)
            elif content_type == 'application/xml':
                processed_content, method, warnings = self._process_xml_content(raw_content)
            elif content_type in ['text/yaml', 'application/yaml']:
                processed_content, method, warnings = self._process_yaml_content(raw_content)
            else:
                # Plain text processing
                processed_content = raw_content
                method = "text_extraction"
                warnings = []
            
            # Estimate tokens
            token_count = self.token_estimator.estimate_text_tokens(processed_content, model)
            
            return ProcessingResult(
                success=True,
                content=processed_content,
                content_type=content_type,
                token_count=token_count,
                processing_method=method,
                warnings=warnings,
                metadata={"original_length": len(raw_content)}
            )
            
        except UnicodeDecodeError as e:
            return ProcessingResult(
                success=False,
                content_type=content_type,
                processing_method="encoding_error",
                warnings=[f"Encoding error: {str(e)}"]
            )
    
    def _process_json_content(self, content: str) -> Tuple[str, str, List[str]]:
        """Process JSON content."""
        warnings = []
        try:
            # Parse JSON to validate and potentially format
            data = json.loads(content)
            
            # For large JSON, provide a summary instead of full content
            if len(content) > 10000:  # 10KB threshold
                if isinstance(data, dict):
                    summary = {
                        "type": "object",
                        "keys": list(data.keys())[:20],  # First 20 keys
                        "total_keys": len(data.keys()) if isinstance(data, dict) else None
                    }
                elif isinstance(data, list):
                    summary = {
                        "type": "array",
                        "length": len(data),
                        "sample_items": data[:5] if len(data) > 5 else data
                    }
                else:
                    summary = {"type": type(data).__name__, "value": str(data)[:500]}
                
                processed_content = f"JSON Summary:\n{json.dumps(summary, indent=2)}\n\nOriginal size: {len(content)} characters"
                warnings.append("Large JSON file summarized for token efficiency")
                method = "json_summary"
            else:
                # Format JSON nicely
                processed_content = json.dumps(data, indent=2, ensure_ascii=False)
                method = "json_formatting"
            
        except json.JSONDecodeError as e:
            processed_content = content  # Return original if parsing fails
            warnings.append(f"JSON parsing failed: {str(e)}")
            method = "json_raw"
        
        return processed_content, method, warnings
    
    def _process_csv_content(self, content: str) -> Tuple[str, str, List[str]]:
        """Process CSV content."""
        warnings = []
        try:
            # Parse CSV
            lines = content.strip().split('\n')
            reader = csv.reader(lines)
            rows = list(reader)
            
            if not rows:
                return content, "csv_empty", ["Empty CSV file"]
            
            # For large CSV files, provide summary + sample
            if len(rows) > 100:  # More than 100 rows
                header = rows[0] if rows else []
                sample_rows = rows[1:6] if len(rows) > 1 else []  # First 5 data rows
                
                processed_content = f"CSV Summary:\n"
                processed_content += f"Columns ({len(header)}): {', '.join(header)}\n"
                processed_content += f"Total rows: {len(rows) - 1}\n\n"
                processed_content += "Sample data:\n"
                
                # Format as table
                if header and sample_rows:
                    processed_content += " | ".join(header) + "\n"
                    processed_content += "-" * (len(" | ".join(header))) + "\n"
                    for row in sample_rows:
                        processed_content += " | ".join(str(cell) for cell in row) + "\n"
                
                warnings.append(f"Large CSV file summarized (showing 5 of {len(rows)-1} rows)")
                method = "csv_summary"
            else:
                # Format as readable table
                if rows:
                    header = rows[0]
                    data_rows = rows[1:]
                    
                    processed_content = f"CSV Data ({len(data_rows)} rows):\n\n"
                    processed_content += " | ".join(header) + "\n"
                    processed_content += "-" * (len(" | ".join(header))) + "\n"
                    
                    for row in data_rows:
                        processed_content += " | ".join(str(cell) for cell in row) + "\n"
                else:
                    processed_content = content
                
                method = "csv_formatting"
            
        except Exception as e:
            processed_content = content
            warnings.append(f"CSV processing failed: {str(e)}")
            method = "csv_raw"
        
        return processed_content, method, warnings
    
    def _process_xml_content(self, content: str) -> Tuple[str, str, List[str]]:
        """Process XML content."""
        warnings = []
        try:
            # Parse XML
            root = ET.fromstring(content)
            
            # For large XML, provide structure summary
            if len(content) > 10000:  # 10KB threshold
                def get_structure(element, max_depth=3, current_depth=0):
                    if current_depth >= max_depth:
                        return f"{element.tag}[...]"
                    
                    children = list(element)
                    if not children:
                        return element.tag
                    
                    child_tags = {}
                    for child in children:
                        tag = child.tag
                        if tag not in child_tags:
                            child_tags[tag] = 0
                        child_tags[tag] += 1
                    
                    structure = element.tag + " {\n"
                    for tag, count in child_tags.items():
                        if count > 1:
                            structure += f"  {tag} (x{count})\n"
                        else:
                            structure += f"  {get_structure(children[0] if children else element, max_depth, current_depth + 1)}\n"
                    structure += "}"
                    
                    return structure
                
                structure = get_structure(root)
                processed_content = f"XML Structure Summary:\n{structure}\n\nOriginal size: {len(content)} characters"
                warnings.append("Large XML file summarized for token efficiency")
                method = "xml_summary"
            else:
                # Return formatted XML (original content is usually fine)
                processed_content = content
                method = "xml_raw"
            
        except ET.ParseError as e:
            processed_content = content
            warnings.append(f"XML parsing failed: {str(e)}")
            method = "xml_raw"
        
        return processed_content, method, warnings
    
    def _process_yaml_content(self, content: str) -> Tuple[str, str, List[str]]:
        """Process YAML content."""
        warnings = []
        try:
            # Try to import yaml
            import yaml
            
            # Parse YAML
            data = yaml.safe_load(content)
            
            # For large YAML, provide summary
            if len(content) > 10000:  # 10KB threshold
                if isinstance(data, dict):
                    summary = {
                        "type": "mapping",
                        "keys": list(data.keys())[:20],
                        "total_keys": len(data.keys())
                    }
                elif isinstance(data, list):
                    summary = {
                        "type": "sequence", 
                        "length": len(data),
                        "sample_items": data[:5] if len(data) > 5 else data
                    }
                else:
                    summary = {"type": type(data).__name__, "value": str(data)[:500]}
                
                processed_content = f"YAML Summary:\n{yaml.dump(summary, default_flow_style=False)}\n\nOriginal size: {len(content)} characters"
                warnings.append("Large YAML file summarized for token efficiency")
                method = "yaml_summary"
            else:
                processed_content = content  # YAML is usually human-readable as-is
                method = "yaml_raw"
            
        except ImportError:
            processed_content = content
            warnings.append("PyYAML not available, processing as plain text")
            method = "yaml_fallback"
        except yaml.YAMLError as e:
            processed_content = content
            warnings.append(f"YAML parsing failed: {str(e)}")
            method = "yaml_raw"
        
        return processed_content, method, warnings
    
    def _process_image_file(self, 
                           file_path: Path, 
                           content_type: str, 
                           model: str) -> ProcessingResult:
        """Process image files (extract metadata, estimate tokens)."""
        try:
            # Check if model supports vision
            if not self.token_estimator.supports_vision(model):
                return ProcessingResult(
                    success=False,
                    content_type=content_type,
                    processing_method="vision_not_supported",
                    warnings=[f"Model {model} does not support vision"]
                )
            
            # Try to get image dimensions and basic info
            try:
                from PIL import Image
                with Image.open(file_path) as img:
                    width, height = img.size
                    mode = img.mode
                    format_name = img.format
                    
                    metadata = {
                        "width": width,
                        "height": height,
                        "mode": mode,
                        "format": format_name,
                        "has_transparency": mode in ('RGBA', 'LA') or 'transparency' in img.info
                    }
            except ImportError:
                metadata = {"error": "PIL not available for image analysis"}
            except Exception as e:
                metadata = {"error": f"Image analysis failed: {str(e)}"}
            
            # Estimate tokens for the image
            token_count = self.token_estimator.estimate_image_tokens(file_path, model)
            
            # Create a text description of the image for content
            content_parts = [f"Image file: {file_path.name}"]
            if "width" in metadata and "height" in metadata:
                content_parts.append(f"Dimensions: {metadata['width']}x{metadata['height']}")
            if "format" in metadata:
                content_parts.append(f"Format: {metadata['format']}")
            
            content = "\n".join(content_parts)
            
            return ProcessingResult(
                success=True,
                content=content,
                content_type=content_type,
                token_count=token_count,
                processing_method="image_metadata",
                metadata=metadata
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                content_type=content_type,
                processing_method="image_error",
                warnings=[f"Image processing error: {str(e)}"]
            )
    
    def create_processed_file(self, 
                            file_path: Union[str, Path], 
                            model: str,
                            content_type: Optional[str] = None) -> ProcessedFile:
        """
        Create a ProcessedFile object from a file.
        
        Args:
            file_path: Path to the file
            model: Model name for token estimation
            content_type: Optional content type override
            
        Returns:
            ProcessedFile object
        """
        result = self.process_file(file_path, model, content_type)
        
        return ProcessedFile(
            original_path=file_path,
            content=result.content,
            content_type=result.content_type,
            token_count=result.token_count,
            processing_method=result.processing_method,
            warnings=result.warnings
        )
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.SUPPORTED_EXTENSIONS.keys())
    
    def is_supported_file(self, file_path: Union[str, Path]) -> bool:
        """Check if a file type is supported."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        return extension in self.SUPPORTED_EXTENSIONS
    
    def estimate_processing_time(self, file_path: Union[str, Path]) -> float:
        """
        Estimate processing time for a file (in seconds).
        
        Args:
            file_path: Path to the file
            
        Returns:
            Estimated processing time in seconds
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return 0.0
        
        file_size = file_path.stat().st_size
        extension = file_path.suffix.lower()
        
        # Base time per MB for different file types
        time_per_mb = {
            '.txt': 0.1,
            '.md': 0.1,
            '.json': 0.3,
            '.csv': 0.5,
            '.xml': 0.4,
            '.yaml': 0.2,
            '.yml': 0.2,
        }
        
        # Image files are faster to process (just metadata)
        if extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
            time_per_mb[extension] = 0.05
        
        base_time = time_per_mb.get(extension, 0.2)  # Default 0.2s per MB
        size_mb = file_size / (1024 * 1024)
        
        return max(0.01, base_time * size_mb)  # Minimum 0.01s
    
    def _process_pdf_file(self, 
                         file_path: Path, 
                         content_type: str, 
                         model: str) -> ProcessingResult:
        """Process PDF files by extracting text content."""
        try:
            # Try PyPDF2 first
            try:
                import PyPDF2
                
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    # Check if PDF is encrypted
                    if pdf_reader.is_encrypted:
                        return ProcessingResult(
                            success=False,
                            content_type=content_type,
                            processing_method="pdf_encrypted",
                            warnings=["PDF is encrypted and cannot be processed"]
                        )
                    
                    # Extract text from all pages
                    text_content = []
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                        except Exception as e:
                            text_content.append(f"--- Page {page_num + 1} ---\n[Error extracting text: {str(e)}]")
                    
                    if not text_content:
                        return ProcessingResult(
                            success=False,
                            content_type=content_type,
                            processing_method="pdf_no_text",
                            warnings=["No extractable text found in PDF"]
                        )
                    
                    full_text = "\n\n".join(text_content)
                    
                    # Estimate tokens
                    token_count = self.token_estimator.estimate_text_tokens(full_text, model)
                    
                    return ProcessingResult(
                        success=True,
                        content=full_text,
                        content_type=content_type,
                        token_count=token_count,
                        processing_method="pdf_pypdf2",
                        metadata={
                            "page_count": len(pdf_reader.pages),
                            "extracted_pages": len(text_content)
                        }
                    )
                    
            except ImportError:
                # Try pdfplumber as fallback
                try:
                    import pdfplumber
                    
                    with pdfplumber.open(file_path) as pdf:
                        text_content = []
                        for page_num, page in enumerate(pdf.pages):
                            try:
                                page_text = page.extract_text()
                                if page_text and page_text.strip():
                                    text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                            except Exception as e:
                                text_content.append(f"--- Page {page_num + 1} ---\n[Error extracting text: {str(e)}]")
                        
                        if not text_content:
                            return ProcessingResult(
                                success=False,
                                content_type=content_type,
                                processing_method="pdf_no_text",
                                warnings=["No extractable text found in PDF"]
                            )
                        
                        full_text = "\n\n".join(text_content)
                        token_count = self.token_estimator.estimate_text_tokens(full_text, model)
                        
                        return ProcessingResult(
                            success=True,
                            content=full_text,
                            content_type=content_type,
                            token_count=token_count,
                            processing_method="pdf_pdfplumber",
                            metadata={
                                "page_count": len(pdf.pages),
                                "extracted_pages": len(text_content)
                            }
                        )
                        
                except ImportError:
                    return ProcessingResult(
                        success=False,
                        content_type=content_type,
                        processing_method="pdf_no_library",
                        warnings=["No PDF processing library available (PyPDF2 or pdfplumber required)"]
                    )
                    
        except Exception as e:
            return ProcessingResult(
                success=False,
                content_type=content_type,
                processing_method="pdf_error",
                warnings=[f"PDF processing error: {str(e)}"]
            )
    
    def _process_word_file(self, 
                          file_path: Path, 
                          content_type: str, 
                          model: str) -> ProcessingResult:
        """Process Word documents (.docx, .doc)."""
        try:
            if content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Process .docx files
                try:
                    import docx
                    
                    doc = docx.Document(file_path)
                    
                    # Extract text from paragraphs
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    
                    # Extract text from tables
                    table_content = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells)
                            if row_text.strip():
                                table_content.append(row_text)
                    
                    # Combine content
                    content_parts = []
                    if paragraphs:
                        content_parts.append("Document Text:\n" + "\n\n".join(paragraphs))
                    if table_content:
                        content_parts.append("Tables:\n" + "\n".join(table_content))
                    
                    if not content_parts:
                        return ProcessingResult(
                            success=False,
                            content_type=content_type,
                            processing_method="docx_no_content",
                            warnings=["No extractable content found in Word document"]
                        )
                    
                    full_text = "\n\n".join(content_parts)
                    token_count = self.token_estimator.estimate_text_tokens(full_text, model)
                    
                    return ProcessingResult(
                        success=True,
                        content=full_text,
                        content_type=content_type,
                        token_count=token_count,
                        processing_method="docx_python_docx",
                        metadata={
                            "paragraph_count": len(paragraphs),
                            "table_count": len(doc.tables)
                        }
                    )
                    
                except ImportError:
                    return ProcessingResult(
                        success=False,
                        content_type=content_type,
                        processing_method="docx_no_library",
                        warnings=["python-docx library required for .docx files"]
                    )
            
            else:
                # .doc files are more complex, return unsupported for now
                return ProcessingResult(
                    success=False,
                    content_type=content_type,
                    processing_method="doc_unsupported",
                    warnings=[".doc files not supported (use .docx format)"]
                )
                
        except Exception as e:
            return ProcessingResult(
                success=False,
                content_type=content_type,
                processing_method="word_error",
                warnings=[f"Word document processing error: {str(e)}"]
            )
    
    def _process_rtf_file(self, 
                         file_path: Path, 
                         content_type: str, 
                         model: str) -> ProcessingResult:
        """Process RTF (Rich Text Format) files."""
        try:
            # Try striprtf library
            try:
                from striprtf.striprtf import rtf_to_text
                
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    rtf_content = f.read()
                
                # Convert RTF to plain text
                plain_text = rtf_to_text(rtf_content)
                
                if not plain_text.strip():
                    return ProcessingResult(
                        success=False,
                        content_type=content_type,
                        processing_method="rtf_no_content",
                        warnings=["No extractable text found in RTF file"]
                    )
                
                token_count = self.token_estimator.estimate_text_tokens(plain_text, model)
                
                return ProcessingResult(
                    success=True,
                    content=plain_text,
                    content_type=content_type,
                    token_count=token_count,
                    processing_method="rtf_striprtf",
                    metadata={"original_rtf_length": len(rtf_content)}
                )
                
            except ImportError:
                # Fallback: try basic RTF parsing (remove RTF codes)
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    rtf_content = f.read()
                
                # Basic RTF cleaning (remove control codes)
                import re
                
                # Remove RTF control words and groups
                text = re.sub(r'\\[a-z]+\d*\s?', '', rtf_content)  # Control words
                text = re.sub(r'[{}]', '', text)  # Braces
                text = re.sub(r'\\[^a-z]', '', text)  # Control symbols
                text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
                text = text.strip()
                
                if not text:
                    return ProcessingResult(
                        success=False,
                        content_type=content_type,
                        processing_method="rtf_no_content",
                        warnings=["No extractable text found in RTF file"]
                    )
                
                token_count = self.token_estimator.estimate_text_tokens(text, model)
                
                return ProcessingResult(
                    success=True,
                    content=text,
                    content_type=content_type,
                    token_count=token_count,
                    processing_method="rtf_basic_parsing",
                    warnings=["Basic RTF parsing used (striprtf library recommended for better results)"],
                    metadata={"original_rtf_length": len(rtf_content)}
                )
                
        except Exception as e:
            return ProcessingResult(
                success=False,
                content_type=content_type,
                processing_method="rtf_error",
                warnings=[f"RTF processing error: {str(e)}"]
            )
    
    def validate_file_integrity(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Validate file integrity and provide diagnostic information.
        
        Args:
            file_path: Path to the file to validate
            
        Returns:
            Dictionary with validation results
        """
        file_path = Path(file_path)
        
        result = {
            "valid": False,
            "exists": False,
            "readable": False,
            "size": 0,
            "content_type": None,
            "issues": []
        }
        
        # Check existence
        if not file_path.exists():
            result["issues"].append("File does not exist")
            return result
        
        result["exists"] = True
        
        # Check if it's a file (not directory)
        if not file_path.is_file():
            result["issues"].append("Path is not a file")
            return result
        
        # Check size
        try:
            result["size"] = file_path.stat().st_size
            if result["size"] == 0:
                result["issues"].append("File is empty")
            elif result["size"] > self.max_file_size:
                result["issues"].append(f"File too large ({result['size']} bytes, max: {self.max_file_size})")
        except Exception as e:
            result["issues"].append(f"Cannot read file size: {str(e)}")
            return result
        
        # Check readability
        try:
            with open(file_path, 'rb') as f:
                f.read(1)  # Try to read first byte
            result["readable"] = True
        except Exception as e:
            result["issues"].append(f"File not readable: {str(e)}")
            return result
        
        # Detect content type
        result["content_type"] = self._detect_content_type(file_path)
        
        # Check if supported
        if not self.is_supported_file(file_path):
            result["issues"].append(f"Unsupported file type: {file_path.suffix}")
        
        # File-type specific validation
        if result["content_type"] == "application/pdf":
            result.update(self._validate_pdf(file_path))
        elif result["content_type"].startswith("image/"):
            result.update(self._validate_image(file_path))
        
        result["valid"] = len(result["issues"]) == 0
        
        return result
    
    def _validate_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Validate PDF file specifically."""
        validation = {"pdf_specific": {}}
        
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                try:
                    reader = PyPDF2.PdfReader(f)
                    validation["pdf_specific"]["page_count"] = len(reader.pages)
                    validation["pdf_specific"]["encrypted"] = reader.is_encrypted
                    
                    if reader.is_encrypted:
                        validation["issues"] = validation.get("issues", [])
                        validation["issues"].append("PDF is encrypted")
                    
                    # Try to read first page
                    if reader.pages:
                        try:
                            reader.pages[0].extract_text()
                            validation["pdf_specific"]["text_extractable"] = True
                        except:
                            validation["pdf_specific"]["text_extractable"] = False
                            validation["issues"] = validation.get("issues", [])
                            validation["issues"].append("PDF text extraction may fail")
                    
                except Exception as e:
                    validation["issues"] = validation.get("issues", [])
                    validation["issues"].append(f"PDF validation error: {str(e)}")
                    
        except ImportError:
            validation["pdf_specific"]["library_available"] = False
            
        return validation
    
    def _validate_image(self, file_path: Path) -> Dict[str, Any]:
        """Validate image file specifically."""
        validation = {"image_specific": {}}
        
        try:
            from PIL import Image
            
            with Image.open(file_path) as img:
                validation["image_specific"]["dimensions"] = img.size
                validation["image_specific"]["mode"] = img.mode
                validation["image_specific"]["format"] = img.format
                
                # Check for potential issues
                width, height = img.size
                if width * height > 50_000_000:  # 50MP
                    validation["issues"] = validation.get("issues", [])
                    validation["issues"].append("Very large image may consume many tokens")
                    
        except ImportError:
            validation["image_specific"]["library_available"] = False
        except Exception as e:
            validation["issues"] = validation.get("issues", [])
            validation["issues"].append(f"Image validation error: {str(e)}")
            
        return validation