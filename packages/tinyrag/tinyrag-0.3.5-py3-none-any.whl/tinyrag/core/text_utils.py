"""
Text extraction and chunking utilities with optimized processing for large files
"""

import os
import time
from typing import List, Union, Iterator
from pathlib import Path


def _extract_pdf_fast(file_path: str) -> str:
    """Fast PDF extraction with fallback options"""
    text = ""
    
    # Try PyMuPDF first (fastest)
    try:
        import fitz
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text
    except ImportError:
        pass
    except Exception as e:
        print(f"PyMuPDF failed: {e}, trying pdfminer...")
    
    # Fallback to pdfminer
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract_text
        return pdfminer_extract_text(file_path)
    except ImportError:
        raise ImportError("PDF processing requires PyMuPDF (pip install PyMuPDF) or pdfminer.six (pip install pdfminer.six)")
    except Exception as e:
        print(f"pdfminer failed: {e}")
        return ""

def _extract_docx_fast(file_path: str) -> str:
    """Fast DOCX extraction with error handling"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = paragraphs + table_text
        return "\n".join(all_text)
        
    except ImportError:
        raise ImportError("python-docx required for DOCX support. Install with: pip install python-docx")
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""


def extract_text(data: Union[str, Path], show_progress: bool = False) -> str:
    """Extract text from file path or return raw string. Supports txt, pdf, docx, csv.
    Optimized for large files with better error handling.
    
    Args:
        data: File path or raw text string
        show_progress: Show progress for large files
    
    Returns:
        Extracted text content
    """
    if isinstance(data, (str, Path)) and os.path.isfile(data):
        file_path = Path(data)
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        extension = file_path.suffix.lower()
        
        if show_progress and file_size_mb > 1:
            print(f"📄 Processing {file_path.name} ({file_size_mb:.1f} MB)...")
        
        start_time = time.time()
        
        try:
            if extension == '.txt':
                # Handle text files with encoding detection
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            text = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # If all encodings fail, read as binary and decode with errors='ignore'
                    with open(file_path, 'rb') as f:
                        text = f.read().decode('utf-8', errors='ignore')
            
            elif extension == '.pdf':
                text = _extract_pdf_fast(str(file_path))
            
            elif extension == '.docx':
                text = _extract_docx_fast(str(file_path))
            
            elif extension == '.csv':
                try:
                    import pandas as pd
                    # Use pandas for faster CSV processing
                    df = pd.read_csv(file_path, encoding='utf-8')
                    text = df.to_string(index=False)
                except ImportError:
                    # Fallback to standard csv
                    import csv
                    with open(file_path, 'r', encoding='utf-8', newline='') as f:
                        reader = csv.reader(f)
                        rows = [" ".join(row) for row in reader]
                        text = "\n".join(rows)
                except Exception as e:
                    print(f"CSV processing failed: {e}")
                    # Try as plain text
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
            
            else:
                # Try to read as plain text for other formats
                encodings = ['utf-8', 'latin-1', 'cp1252']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            text = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # Binary fallback
                    with open(file_path, 'rb') as f:
                        text = f.read().decode('utf-8', errors='ignore')
            
            # Show processing time for large files
            if show_progress and file_size_mb > 1:
                elapsed = time.time() - start_time
                print(f"✓ Processed in {elapsed:.1f}s")
            
            # Validate that we got meaningful content
            if not text or not text.strip():
                print(f"⚠ Warning: No content extracted from {file_path.name}")
                return ""
            
            return text
            
        except Exception as e:
            print(f"⚠ Error processing {file_path.name}: {e}")
            return ""
    
    # Return as raw string if not a file path
    return str(data) if data else ""


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Split text into overlapping chunks with smart boundary detection
    
    Args:
        text: Input text to chunk
        chunk_size: Maximum size of each chunk
        overlap: Number of characters to overlap between chunks
    
    Returns:
        List of text chunks
    """
    if not text or not text.strip():
        return []
    
    text = text.strip()
    
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Don't go beyond the text
        if end >= len(text):
            chunk = text[start:].strip()
            if chunk:
                chunks.append(chunk)
            break
        
        # Try to break at natural boundaries
        best_end = end
        
        # Look for paragraph break (double newline)
        para_break = text.rfind('\n\n', start, end)
        if para_break > start + chunk_size // 2:  # Only if it's not too early
            best_end = para_break + 2
        else:
            # Look for sentence boundary
            sentence_breaks = ['. ', '! ', '? ', '.\n', '!\n', '?\n']
            for break_char in sentence_breaks:
                sentence_end = text.rfind(break_char, start + chunk_size // 2, end)
                if sentence_end > start:
                    best_end = sentence_end + len(break_char)
                    break
            else:
                # Look for word boundary
                word_end = text.rfind(' ', start + chunk_size // 2, end)
                if word_end > start:
                    best_end = word_end
                else:
                    # Look for any whitespace
                    whitespace_end = text.rfind('\n', start, end)
                    if whitespace_end > start:
                        best_end = whitespace_end + 1
        
        chunk = text[start:best_end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Calculate next start position with overlap
        start = best_end - overlap
        if start <= 0 or start >= len(text):
            break
        
        # Ensure we're making progress
        if start >= best_end:
            start = best_end
    
    # Remove duplicate chunks and very short chunks
    unique_chunks = []
    seen = set()
    
    for chunk in chunks:
        if len(chunk) >= 20 and chunk not in seen:  # Minimum chunk size
            unique_chunks.append(chunk)
            seen.add(chunk)
    
    return unique_chunks